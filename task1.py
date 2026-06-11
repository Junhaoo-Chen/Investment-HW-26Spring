#!/usr/bin/env python3
"""Task 1: Markowitz Portfolio Optimization for 10 SW Industry Indices (2015-2025)."""

import os
import numpy as np
import pandas as pd
from scipy.optimize import minimize
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from matplotlib.font_manager import FontProperties

_CN_FONT = FontProperties(fname='/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc')
plt.rcParams['axes.unicode_minus'] = False
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────────────────────────
DATA_DIR = '/cpfs02/data/shared/Group-m6/yanhai.cjh/Misc/investment_hw/数据文件/'
OUTPUT_DIR = '/cpfs02/data/shared/Group-m6/yanhai.cjh/Misc/investment_hw/output/'

INDEX_NAMES = {
    801010: '农林牧渔', 801050: '有色金属', 801080: '电子',
    801110: '家用电器', 801120: '食品饮料', 801150: '医药生物',
    801160: '公用事业', 801180: '房地产',   801780: '银行',
    801890: '机械设备',
}
INDEX_NAMES_EN = {
    801010: 'Agriculture', 801050: 'Non-Ferrous Metals', 801080: 'Electronics',
    801110: 'Home Appliances', 801120: 'Food & Beverage', 801150: 'Pharma & Bio',
    801160: 'Utilities', 801180: 'Real Estate', 801780: 'Banking',
    801890: 'Machinery',
}
INDEX_ORDER = [801010, 801050, 801080, 801110, 801120, 801150, 801160, 801180, 801780, 801890]

TRADING_DAYS = 252
A_RISK_AVERSION = 3


# ══════════════════════════════════════════════════════════════════════════════
# 1. DATA LOADING
# ══════════════════════════════════════════════════════════════════════════════

def load_industry_data():
    files = ['20150101-20171231.xlsx', '20180101-20211231.xlsx', '20220101-20251231.xlsx']
    dfs = []
    for f in files:
        df = pd.read_excel(DATA_DIR + f, skiprows=[1, 2])
        dfs.append(df)
    industry = pd.concat(dfs, ignore_index=True)
    returns_pct = industry.pivot(index='Idxtrd01', columns='Indexcd', values='Idxtrd08')
    returns_pct = returns_pct[INDEX_ORDER].sort_index()
    returns_decimal = returns_pct / 100.0
    assert returns_decimal.shape == (2674, 10), f"Expected (2674,10), got {returns_decimal.shape}"
    assert returns_decimal.isna().sum().sum() == 0, "NaN values found"
    return returns_decimal


def load_shibor(trading_dates):
    shibor = pd.read_excel(DATA_DIR + 'SHIBOR.xlsx', header=0)
    shibor = shibor.rename(columns={
        '交易日期_TrdDt': 'date',
        '加权价(%)_WghAvgPr': 'rate_annual_pct',
    })
    shibor = shibor[shibor['date'].isin(trading_dates)]
    shibor = shibor.set_index('date').sort_index()
    shibor['rf_daily'] = shibor['rate_annual_pct'] / (360 * 100)
    return shibor['rf_daily']


# ══════════════════════════════════════════════════════════════════════════════
# 2. DESCRIPTIVE STATISTICS
# ══════════════════════════════════════════════════════════════════════════════

def compute_descriptive_stats(returns):
    daily_mean = returns.mean()
    daily_std = returns.std(ddof=1)
    annual_mean = daily_mean * TRADING_DAYS
    annual_std = daily_std * np.sqrt(TRADING_DAYS)

    stats = pd.DataFrame({
        'Daily Mean Return': daily_mean,
        'Daily Std Dev': daily_std,
        'Annualized Mean Return': annual_mean,
        'Annualized Std Dev': annual_std,
    })
    stats.index = [f"{INDEX_NAMES[c]} ({c})" for c in stats.index]
    return stats


def compute_correlation(returns):
    corr = returns.corr()
    corr.index = [INDEX_NAMES[c] for c in corr.index]
    corr.columns = [INDEX_NAMES[c] for c in corr.columns]
    return corr


# ══════════════════════════════════════════════════════════════════════════════
# 3. MARKOWITZ STEP 1: EFFICIENT FRONTIER (RISKY ASSETS ONLY)
# ══════════════════════════════════════════════════════════════════════════════

def compute_frontier_analytical(mu, Sigma, tang_mu=None):
    n = len(mu)
    ones = np.ones(n)
    Sigma_inv = np.linalg.inv(Sigma)

    A_val = ones @ Sigma_inv @ mu
    B_val = mu @ Sigma_inv @ mu
    C_val = ones @ Sigma_inv @ ones
    D_val = B_val * C_val - A_val ** 2

    w_gmv = Sigma_inv @ ones / C_val
    mu_gmv = A_val / C_val
    sigma_gmv = np.sqrt(1.0 / C_val)

    upper_limit = mu.max() + 0.0008
    if tang_mu is not None:
        upper_limit = max(upper_limit, tang_mu * 1.15)
    mu_range = np.linspace(mu_gmv - 0.0005, upper_limit, 800)
    sigma_range = np.sqrt((C_val * mu_range**2 - 2 * A_val * mu_range + B_val) / D_val)

    upper_mask = mu_range >= mu_gmv
    return {
        'w_gmv': w_gmv, 'mu_gmv': mu_gmv, 'sigma_gmv': sigma_gmv,
        'A': A_val, 'B': B_val, 'C': C_val, 'D': D_val,
        'mu_range': mu_range, 'sigma_range': sigma_range,
        'upper_mask': upper_mask,
    }


# ══════════════════════════════════════════════════════════════════════════════
# 4. MARKOWITZ STEP 2: TANGENCY PORTFOLIO & CML
# ══════════════════════════════════════════════════════════════════════════════

def compute_tangency(mu, Sigma, rf):
    ones = np.ones(len(mu))
    Sigma_inv = np.linalg.inv(Sigma)
    excess_mu = mu - rf * ones
    w_tang = Sigma_inv @ excess_mu / (ones @ Sigma_inv @ excess_mu)

    mu_tang = w_tang @ mu
    sigma_tang = np.sqrt(w_tang @ Sigma @ w_tang)
    sharpe = (mu_tang - rf) / sigma_tang

    assert abs(w_tang.sum() - 1.0) < 1e-10, f"Tangency weights sum to {w_tang.sum()}"

    return {
        'w': w_tang, 'mu': mu_tang, 'sigma': sigma_tang, 'sharpe': sharpe,
    }


def compute_cml(rf, mu_tang, sigma_tang, sigma_max):
    sigma_cml = np.linspace(0, sigma_max, 300)
    mu_cml = rf + (mu_tang - rf) / sigma_tang * sigma_cml
    return sigma_cml, mu_cml


# ══════════════════════════════════════════════════════════════════════════════
# 5. MARKOWITZ STEP 3: COMPLETE PORTFOLIO
# ══════════════════════════════════════════════════════════════════════════════

def compute_complete_portfolio(mu_tang, sigma_tang, rf, A):
    y_star = (mu_tang - rf) / (A * sigma_tang ** 2)
    mu_complete = rf + y_star * (mu_tang - rf)
    sigma_complete = abs(y_star) * sigma_tang
    return y_star, mu_complete, sigma_complete


# ══════════════════════════════════════════════════════════════════════════════
# 6. PLOTTING
# ══════════════════════════════════════════════════════════════════════════════

def plot_efficient_frontier(frontier, tangency, rf, mu, Sigma, y_star, mu_complete, sigma_complete, lang='en'):
    fig, ax = plt.subplots(figsize=(10, 7))
    fp = _CN_FONT if lang == 'cn' else None

    ann_factor_mu = TRADING_DAYS * 100
    ann_factor_sigma = np.sqrt(TRADING_DAYS) * 100

    f_sigma = frontier['sigma_range'] * ann_factor_sigma
    f_mu = frontier['mu_range'] * ann_factor_mu

    upper = frontier['upper_mask']
    ax.plot(f_sigma[~upper], f_mu[~upper], 'b--', alpha=0.4, linewidth=1)
    lbl_ef = 'Efficient Frontier' if lang == 'en' else '有效前沿'
    ax.plot(f_sigma[upper], f_mu[upper], 'b-', linewidth=2, label=lbl_ef)

    sigma_cml, mu_cml = compute_cml(rf, tangency['mu'], tangency['sigma'],
                                     frontier['sigma_range'].max() * 1.3)
    lbl_cml = 'Capital Market Line (CML)' if lang == 'en' else '资本市场线 (CML)'
    ax.plot(sigma_cml * ann_factor_sigma, mu_cml * ann_factor_mu,
            'r-', linewidth=2, label=lbl_cml)

    lbl_gmv = 'Global Min-Variance Portfolio' if lang == 'en' else '全局最小方差组合'
    ax.plot(frontier['sigma_gmv'] * ann_factor_sigma, frontier['mu_gmv'] * ann_factor_mu,
            'gs', markersize=10, label=lbl_gmv, zorder=5)

    lbl_tang = 'Tangency Portfolio' if lang == 'en' else '切线组合 (最优风险组合)'
    ax.plot(tangency['sigma'] * ann_factor_sigma, tangency['mu'] * ann_factor_mu,
            'r*', markersize=18, label=lbl_tang, zorder=5)

    lbl_comp = f'Complete Portfolio (y*={y_star:.4f})' if lang == 'en' else f'完整组合 (y*={y_star:.4f})'
    ax.plot(sigma_complete * ann_factor_sigma, mu_complete * ann_factor_mu,
            'mo', markersize=10, label=lbl_comp, zorder=5)

    for i, code in enumerate(INDEX_ORDER):
        sigma_i = np.sqrt(Sigma[i, i]) * ann_factor_sigma
        mu_i = mu[i] * ann_factor_mu
        ax.plot(sigma_i, mu_i, 'k^', markersize=6, zorder=4)
        ax.annotate(INDEX_NAMES[code], (sigma_i, mu_i),
                    textcoords="offset points", xytext=(5, 5), fontsize=7,
                    fontproperties=_CN_FONT)

    xlabel = 'Annualized Standard Deviation (%)' if lang == 'en' else '年化标准差 (%)'
    ylabel = 'Annualized Expected Return (%)' if lang == 'en' else '年化期望收益率 (%)'
    title = ('Efficient Frontier, CML, and Optimal Portfolios' if lang == 'en'
             else '有效前沿、资本市场线与最优组合')
    ax.set_xlabel(xlabel, fontsize=12, fontproperties=fp)
    ax.set_ylabel(ylabel, fontsize=12, fontproperties=fp)
    ax.set_title(title, fontsize=14, fontproperties=fp)
    ax.legend(fontsize=9, loc='upper left', prop=_CN_FONT if lang == 'cn' else None)
    ax.grid(True, alpha=0.3)

    suffix = '' if lang == 'en' else '_cn'
    path = OUTPUT_DIR + f'efficient_frontier{suffix}.png'
    fig.savefig(path, dpi=300, bbox_inches='tight')
    plt.close(fig)
    print(f"Saved {path}")
    return path


def plot_weights(w_tang, lang='en'):
    fig, ax = plt.subplots(figsize=(10, 5))
    fp = _CN_FONT if lang == 'cn' else None
    names = [INDEX_NAMES[c] for c in INDEX_ORDER]
    colors = ['#2196F3' if w >= 0 else '#F44336' for w in w_tang]

    bars = ax.bar(names, w_tang * 100, color=colors, edgecolor='black', linewidth=0.5)
    ax.axhline(y=0, color='black', linewidth=0.5)
    ylabel = 'Weight (%)' if lang == 'en' else '权重 (%)'
    title = 'Tangency Portfolio Weights' if lang == 'en' else '切线组合（最优风险组合）权重'
    ax.set_ylabel(ylabel, fontsize=12, fontproperties=fp)
    ax.set_title(title, fontsize=14, fontproperties=fp)
    ax.set_xticklabels(names, rotation=30, ha='right', fontsize=9, fontproperties=_CN_FONT)

    for bar, w in zip(bars, w_tang):
        ypos = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2., ypos + (1 if ypos >= 0 else -2),
                f'{w*100:.2f}%', ha='center', va='bottom' if ypos >= 0 else 'top', fontsize=8)

    ax.grid(axis='y', alpha=0.3)
    suffix = '' if lang == 'en' else '_cn'
    path = OUTPUT_DIR + f'portfolio_weights{suffix}.png'
    fig.savefig(path, dpi=300, bbox_inches='tight')
    plt.close(fig)
    print(f"Saved {path}")
    return path


# ══════════════════════════════════════════════════════════════════════════════
# 7. WORD DOCUMENT GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_font(cell, text, font_name='Times New Roman', size=9, bold=False, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def add_table_to_doc(doc, df, title=None, fmt='.4f', index_label='Index'):
    if title:
        doc.add_paragraph(title, style='Heading 3')

    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns) + 1)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_cell_font(table.rows[0].cells[0], index_label, bold=True, size=8)
    for j, col in enumerate(df.columns):
        set_cell_font(table.rows[0].cells[j + 1], str(col), bold=True, size=8)

    for i, (idx, row) in enumerate(df.iterrows()):
        set_cell_font(table.rows[i + 1].cells[0], str(idx), size=8)
        for j, val in enumerate(row):
            if isinstance(val, float):
                set_cell_font(table.rows[i + 1].cells[j + 1], format(val, fmt), size=8,
                              alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                set_cell_font(table.rows[i + 1].cells[j + 1], str(val), size=8)
    return table


def add_code_block(doc, code_text, title=None):
    if title:
        doc.add_paragraph(title, style='Heading 3')
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def generate_report(stats, corr, frontier, tangency, rf_avg, rf_annual_avg_pct,
                    y_star, mu_complete, sigma_complete,
                    returns, ef_path, wt_path, lang='en'):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── Title ──
    if lang == 'en':
        title = doc.add_heading('Investment Portfolio Optimization Report', level=0)
        doc.add_paragraph('Project of Investment, 2026')
    else:
        title = doc.add_heading('投资组合优化报告', level=0)
        doc.add_paragraph('2026年投资学课程项目')

    # ── Section 1: Data Description ──
    if lang == 'en':
        doc.add_heading('1. Data Description and Assumptions', level=1)

        doc.add_heading('1.1 Data Description', level=2)
        doc.add_paragraph(
            'This study examines 10 Shenwan (SW) Level-1 industry indices over the period '
            'from January 5, 2015 to December 31, 2025, covering 2,674 trading days. '
            'The 10 industries are: 农林牧渔 (Agriculture, 801010), 有色金属 (Non-Ferrous Metals, 801050), '
            '电子 (Electronics, 801080), 家用电器 (Home Appliances, 801110), 食品饮料 (Food & Beverage, 801120), '
            '医药生物 (Pharma & Bio, 801150), 公用事业 (Utilities, 801160), 房地产 (Real Estate, 801180), '
            '银行 (Banking, 801780), and 机械设备 (Machinery, 801890).'
        )
        doc.add_paragraph(
            'Daily closing index levels and daily returns (in percentage form) were downloaded '
            'from the CSMAR (China Stock Market & Accounting Research) database. '
            'The risk-free rate is the overnight SHIBOR (Shanghai Interbank Offered Rate), '
            'downloaded from the RESSET database for the same period.'
        )

        doc.add_heading('1.2 Risk-Free Rate Conversion', level=2)
        doc.add_paragraph(
            'The overnight SHIBOR is quoted as an annualized percentage rate using the Actual/360 '
            'day-count convention, which is the standard for the Chinese interbank market. '
            'To convert to a daily rate in decimal form:'
        )
        doc.add_paragraph(
            '    r_daily = r_annual(%) / (360 x 100)',
            style='No Spacing'
        )
        doc.add_paragraph(
            f'The average overnight SHIBOR over the sample period is {rf_annual_avg_pct:.4f}% per annum, '
            f'corresponding to an average daily risk-free rate of {rf_avg*100:.6f}% '
            f'(or {rf_avg:.8f} in decimal).'
        )

        doc.add_heading('1.3 Assumptions', level=2)
        doc.add_paragraph(
            '(1) Investors are mean-variance optimizers following the Markowitz framework. '
            '(2) Short-selling is allowed: portfolio weights are unconstrained and may be negative. '
            '(3) The risk aversion coefficient is A = 3, with utility function U = E(r) - (1/2) x A x sigma^2. '
            '(4) Returns are independently and identically distributed (i.i.d.) over the sample period. '
            '(5) This is a single-period investment model with no transaction costs or taxes. '
            '(6) All assets are perfectly divisible and liquid. '
            '(7) Investors can borrow and lend at the risk-free rate (SHIBOR overnight).'
        )
    else:
        doc.add_heading('1. 数据说明与假设', level=1)

        doc.add_heading('1.1 数据说明', level=2)
        doc.add_paragraph(
            '本研究选取了2015年1月5日至2025年12月31日共2,674个交易日的10个申万一级行业指数。'
            '10个行业分别为：农林牧渔 (801010)、有色金属 (801050)、电子 (801080)、'
            '家用电器 (801110)、食品饮料 (801120)、医药生物 (801150)、公用事业 (801160)、'
            '房地产 (801180)、银行 (801780) 和 机械设备 (801890)。'
        )
        doc.add_paragraph(
            '日收盘指数和日收益率（百分比形式）数据来源于国泰安 (CSMAR) 数据库。'
            '无风险利率采用隔夜SHIBOR（上海银行间同业拆放利率），数据来源于锐思 (RESSET) 数据库。'
        )

        doc.add_heading('1.2 无风险利率转换', level=2)
        doc.add_paragraph(
            'SHIBOR隔夜利率以年化百分比利率形式报价，采用 Actual/360 计息惯例（中国银行间市场标准）。'
            '日度利率转换公式为：'
        )
        doc.add_paragraph(
            '    r_daily = r_annual(%) / (360 x 100)',
            style='No Spacing'
        )
        doc.add_paragraph(
            f'样本期内隔夜SHIBOR的平均值为年化 {rf_annual_avg_pct:.4f}%，'
            f'对应日均无风险利率为 {rf_avg*100:.6f}%（即小数形式 {rf_avg:.8f}）。'
        )

        doc.add_heading('1.3 假设', level=2)
        doc.add_paragraph(
            '(1) 投资者为均值-方差优化者，遵循 Markowitz 框架。'
            '(2) 允许卖空：组合权重不受约束，可以为负。'
            '(3) 风险厌恶系数 A = 3，效用函数为 U = E(r) - (1/2) x A x sigma^2。'
            '(4) 收益率在样本期内独立同分布 (i.i.d.)。'
            '(5) 单期投资模型，不考虑交易成本和税收。'
            '(6) 所有资产完全可分且流动性充足。'
            '(7) 投资者可以按无风险利率（SHIBOR隔夜）自由借贷。'
        )

    # ── Table 1: Descriptive Statistics ──
    stats_display = stats.copy()
    stats_display['Daily Mean Return'] = stats_display['Daily Mean Return'] * 100
    stats_display['Daily Std Dev'] = stats_display['Daily Std Dev'] * 100
    stats_display['Annualized Mean Return'] = stats_display['Annualized Mean Return'] * 100
    stats_display['Annualized Std Dev'] = stats_display['Annualized Std Dev'] * 100

    if lang == 'cn':
        stats_display.columns = ['日均收益率(%)', '日标准差(%)', '年化收益率(%)', '年化标准差(%)']
        add_table_to_doc(doc, stats_display, 'Table 1: 描述性统计', fmt='.4f', index_label='行业')
    else:
        add_table_to_doc(doc, stats_display, 'Table 1: Descriptive Statistics of Daily Returns', fmt='.4f')

    # ── Table 2: Correlation Matrix ──
    if lang == 'cn':
        add_table_to_doc(doc, corr, 'Table 2: 行业指数收益率相关系数矩阵', fmt='.4f', index_label='')
    else:
        add_table_to_doc(doc, corr, 'Table 2: Correlation Matrix of Daily Returns', fmt='.4f', index_label='')

    # ── Section 2: Step 1 — Efficient Frontier ──
    if lang == 'en':
        doc.add_heading('2. Markowitz Step 1: Efficient Frontier of Risky Assets', level=1)
        doc.add_paragraph(
            'In the first step, we consider only the 10 risky industry indices (no risk-free asset). '
            'We compute the Global Minimum-Variance Portfolio (GMV) and the minimum-variance frontier.'
        )
        doc.add_paragraph(
            'Let mu denote the (10x1) vector of mean daily returns and Sigma the (10x10) covariance matrix. '
            'The minimum-variance frontier is derived analytically using the two-fund theorem. Define:'
        )
        doc.add_paragraph(
            '    A = 1\' * Sigma^(-1) * mu\n'
            '    B = mu\' * Sigma^(-1) * mu\n'
            '    C = 1\' * Sigma^(-1) * 1\n'
            '    D = B*C - A^2',
            style='No Spacing'
        )
        doc.add_paragraph(
            'The GMV portfolio weights are: w_gmv = Sigma^(-1) * 1 / C'
        )
        doc.add_paragraph(
            f'GMV portfolio: E(r) = {frontier["mu_gmv"]*TRADING_DAYS*100:.4f}% (annualized), '
            f'sigma = {frontier["sigma_gmv"]*np.sqrt(TRADING_DAYS)*100:.4f}% (annualized).'
        )
        doc.add_paragraph(
            'For any target return E, the minimum variance is: sigma^2 = (C*E^2 - 2*A*E + B) / D. '
            'The upper branch of this hyperbola (above the GMV) constitutes the efficient frontier.'
        )
    else:
        doc.add_heading('2. Markowitz 第一步：纯风险资产有效前沿', level=1)
        doc.add_paragraph(
            '第一步仅考虑10个风险行业指数（不含无风险资产），计算全局最小方差组合 (GMV) 和最小方差前沿。'
        )
        doc.add_paragraph(
            '设 mu 为 (10x1) 的日均收益率向量，Sigma 为 (10x10) 的协方差矩阵。'
            '利用两基金分离定理解析求解最小方差前沿。定义：'
        )
        doc.add_paragraph(
            '    A = 1\' * Sigma^(-1) * mu\n'
            '    B = mu\' * Sigma^(-1) * mu\n'
            '    C = 1\' * Sigma^(-1) * 1\n'
            '    D = B*C - A^2',
            style='No Spacing'
        )
        doc.add_paragraph(
            'GMV 组合权重为：w_gmv = Sigma^(-1) * 1 / C'
        )
        doc.add_paragraph(
            f'GMV 组合：年化期望收益率 = {frontier["mu_gmv"]*TRADING_DAYS*100:.4f}%, '
            f'年化标准差 = {frontier["sigma_gmv"]*np.sqrt(TRADING_DAYS)*100:.4f}%。'
        )
        doc.add_paragraph(
            '对于任意目标收益率 E，最小方差为：sigma^2 = (C*E^2 - 2*A*E + B) / D。'
            '该双曲线 GMV 以上的部分即为有效前沿。'
        )

    # Key code for Step 1
    code_step1 = (
        "# Step 1: Efficient Frontier (risky assets only)\n"
        "mu = returns.mean().values          # (10,) daily mean returns\n"
        "Sigma = returns.cov().values         # (10,10) covariance matrix\n"
        "Sigma_inv = np.linalg.inv(Sigma)\n"
        "ones = np.ones(10)\n"
        "A = ones @ Sigma_inv @ mu\n"
        "B = mu @ Sigma_inv @ mu\n"
        "C = ones @ Sigma_inv @ ones\n"
        "D = B * C - A**2\n"
        "w_gmv = Sigma_inv @ ones / C        # GMV portfolio weights\n"
        "# Frontier: sigma^2 = (C*E^2 - 2*A*E + B) / D"
    )
    if lang == 'en':
        add_code_block(doc, code_step1, 'Key Code: Step 1')
    else:
        add_code_block(doc, code_step1, '核心代码：第一步')

    # ── Section 3: Step 2 — Tangency Portfolio & CML ──
    if lang == 'en':
        doc.add_heading('3. Markowitz Step 2: Tangency Portfolio & CML', level=1)
        doc.add_paragraph(
            'In the second step, we introduce the risk-free asset (overnight SHIBOR) and find the tangency '
            'portfolio — the point on the efficient frontier that maximizes the Sharpe ratio. '
            'The tangency portfolio is the optimal risky portfolio.'
        )
        doc.add_paragraph(
            'With short-selling allowed, the analytical solution is:'
        )
        doc.add_paragraph(
            '    w* = Sigma^(-1) * (mu - r_f * 1) / [1\' * Sigma^(-1) * (mu - r_f * 1)]',
            style='No Spacing'
        )
    else:
        doc.add_heading('3. Markowitz 第二步：切线组合与资本市场线 (CML)', level=1)
        doc.add_paragraph(
            '第二步引入无风险资产（SHIBOR隔夜利率），寻找有效前沿上夏普比率最大化的切线组合（即最优风险组合）。'
        )
        doc.add_paragraph(
            '在允许卖空条件下，解析解为：'
        )
        doc.add_paragraph(
            '    w* = Sigma^(-1) * (mu - r_f * 1) / [1\' * Sigma^(-1) * (mu - r_f * 1)]',
            style='No Spacing'
        )

    # Tangency portfolio weights table
    w_tang_df = pd.DataFrame({
        ('Weight (%)' if lang == 'en' else '权重(%)'): tangency['w'] * 100
    }, index=[f"{INDEX_NAMES[c]}" for c in INDEX_ORDER])
    add_table_to_doc(doc, w_tang_df,
                     'Table 3: Tangency Portfolio Weights' if lang == 'en' else 'Table 3: 切线组合权重',
                     fmt='.4f', index_label=('Industry' if lang == 'en' else '行业'))

    if lang == 'en':
        doc.add_paragraph(
            f'Tangency portfolio (annualized): E(r) = {tangency["mu"]*TRADING_DAYS*100:.4f}%, '
            f'sigma = {tangency["sigma"]*np.sqrt(TRADING_DAYS)*100:.4f}%, '
            f'Sharpe ratio = {tangency["sharpe"]*np.sqrt(TRADING_DAYS):.4f}.'
        )
        doc.add_paragraph(
            'The Capital Market Line (CML) passes through the risk-free point (0, r_f) and the '
            'tangency portfolio (sigma_p, E(r_p)):'
        )
        doc.add_paragraph(
            '    E(r_c) = r_f + [(E(r_p) - r_f) / sigma_p] * sigma_c',
            style='No Spacing'
        )
        doc.add_paragraph(
            f'The slope of the CML (reward-to-variability ratio) is the annualized Sharpe ratio: '
            f'{tangency["sharpe"]*np.sqrt(TRADING_DAYS):.4f}.'
        )
    else:
        doc.add_paragraph(
            f'切线组合（年化）：期望收益率 = {tangency["mu"]*TRADING_DAYS*100:.4f}%, '
            f'标准差 = {tangency["sigma"]*np.sqrt(TRADING_DAYS)*100:.4f}%, '
            f'夏普比率 = {tangency["sharpe"]*np.sqrt(TRADING_DAYS):.4f}。'
        )
        doc.add_paragraph(
            '资本市场线 (CML) 过无风险利率点 (0, r_f) 和切线组合点 (sigma_p, E(r_p))：'
        )
        doc.add_paragraph(
            '    E(r_c) = r_f + [(E(r_p) - r_f) / sigma_p] * sigma_c',
            style='No Spacing'
        )
        doc.add_paragraph(
            f'CML 的斜率（报酬-波动比率）即为年化夏普比率：{tangency["sharpe"]*np.sqrt(TRADING_DAYS):.4f}。'
        )

    code_step2 = (
        "# Step 2: Tangency Portfolio & CML\n"
        "rf = shibor_daily.mean()             # average daily risk-free rate\n"
        "excess_mu = mu - rf * np.ones(10)\n"
        "w_tang = Sigma_inv @ excess_mu / (ones @ Sigma_inv @ excess_mu)\n"
        "mu_tang = w_tang @ mu\n"
        "sigma_tang = np.sqrt(w_tang @ Sigma @ w_tang)\n"
        "sharpe = (mu_tang - rf) / sigma_tang  # daily Sharpe ratio\n"
        "# CML: E(r_c) = rf + sharpe * sigma_c"
    )
    if lang == 'en':
        add_code_block(doc, code_step2, 'Key Code: Step 2')
    else:
        add_code_block(doc, code_step2, '核心代码：第二步')

    # ── Section 4: Step 3 — Complete Portfolio ──
    if lang == 'en':
        doc.add_heading('4. Markowitz Step 3: Complete Portfolio (Utility Maximization)', level=1)
        doc.add_paragraph(
            'In the third step, we use the investor\'s utility function to determine the optimal '
            'allocation between the tangency portfolio and the risk-free asset.'
        )
        doc.add_paragraph(
            'The investor maximizes: U = E(r_c) - (1/2) * A * sigma_c^2'
        )
        doc.add_paragraph(
            'Substituting E(r_c) = r_f + y*(E(r_p) - r_f) and sigma_c = y*sigma_p:'
        )
        doc.add_paragraph(
            '    U = r_f + y*(E(r_p) - r_f) - (1/2)*A*y^2*sigma_p^2',
            style='No Spacing'
        )
        doc.add_paragraph(
            'Taking the First-Order Condition (FOC): dU/dy = 0'
        )
        doc.add_paragraph(
            '    dU/dy = (E(r_p) - r_f) - A*y*sigma_p^2 = 0',
            style='No Spacing'
        )
        doc.add_paragraph(
            'Solving for y*:'
        )
        doc.add_paragraph(
            '    y* = (E(r_p) - r_f) / (A * sigma_p^2)',
            style='No Spacing'
        )
        doc.add_paragraph(
            f'With A = {A_RISK_AVERSION}, E(r_p) = {tangency["mu"]*100:.6f}% (daily), '
            f'r_f = {rf_avg*100:.6f}% (daily), '
            f'sigma_p = {tangency["sigma"]*100:.6f}% (daily):'
        )
        doc.add_paragraph(
            f'    y* = ({tangency["mu"]*100:.6f}% - {rf_avg*100:.6f}%) / '
            f'({A_RISK_AVERSION} x ({tangency["sigma"]*100:.6f}%)^2) = {y_star:.4f}',
            style='No Spacing'
        )

        if y_star > 1:
            doc.add_paragraph(
                f'Since y* = {y_star:.4f} > 1, the investor borrows {(y_star-1)*100:.2f}% '
                f'of wealth at the risk-free rate and invests {y_star*100:.2f}% in the tangency portfolio.'
            )
        else:
            doc.add_paragraph(
                f'Since y* = {y_star:.4f}, the investor allocates {y_star*100:.2f}% of wealth '
                f'to the tangency portfolio and {(1-y_star)*100:.2f}% to the risk-free asset.'
            )

        doc.add_paragraph(
            f'Complete portfolio (annualized): E(r_c) = {mu_complete*TRADING_DAYS*100:.4f}%, '
            f'sigma_c = {sigma_complete*np.sqrt(TRADING_DAYS)*100:.4f}%.'
        )
    else:
        doc.add_heading('4. Markowitz 第三步：完整组合（效用最大化）', level=1)
        doc.add_paragraph(
            '第三步利用投资者的效用函数，确定在切线组合与无风险资产之间的最优配置比例。'
        )
        doc.add_paragraph(
            '投资者最大化效用函数：U = E(r_c) - (1/2) * A * sigma_c^2'
        )
        doc.add_paragraph(
            '代入 E(r_c) = r_f + y*(E(r_p) - r_f) 和 sigma_c = y*sigma_p：'
        )
        doc.add_paragraph(
            '    U = r_f + y*(E(r_p) - r_f) - (1/2)*A*y^2*sigma_p^2',
            style='No Spacing'
        )
        doc.add_paragraph(
            '一阶条件 (FOC)：dU/dy = 0'
        )
        doc.add_paragraph(
            '    dU/dy = (E(r_p) - r_f) - A*y*sigma_p^2 = 0',
            style='No Spacing'
        )
        doc.add_paragraph(
            '解得 y*：'
        )
        doc.add_paragraph(
            '    y* = (E(r_p) - r_f) / (A * sigma_p^2)',
            style='No Spacing'
        )
        doc.add_paragraph(
            f'其中 A = {A_RISK_AVERSION}，E(r_p) = {tangency["mu"]*100:.6f}%（日度），'
            f'r_f = {rf_avg*100:.6f}%（日度），'
            f'sigma_p = {tangency["sigma"]*100:.6f}%（日度）：'
        )
        doc.add_paragraph(
            f'    y* = ({tangency["mu"]*100:.6f}% - {rf_avg*100:.6f}%) / '
            f'({A_RISK_AVERSION} x ({tangency["sigma"]*100:.6f}%)^2) = {y_star:.4f}',
            style='No Spacing'
        )

        if y_star > 1:
            doc.add_paragraph(
                f'由于 y* = {y_star:.4f} > 1，投资者需要以无风险利率借入 {(y_star-1)*100:.2f}% '
                f'的资金，将总财富的 {y_star*100:.2f}% 投资于切线组合（杠杆投资）。'
            )
        else:
            doc.add_paragraph(
                f'由于 y* = {y_star:.4f}，投资者将 {y_star*100:.2f}% 的财富投资于切线组合，'
                f'{(1-y_star)*100:.2f}% 投资于无风险资产。'
            )

        doc.add_paragraph(
            f'完整组合（年化）：期望收益率 = {mu_complete*TRADING_DAYS*100:.4f}%, '
            f'标准差 = {sigma_complete*np.sqrt(TRADING_DAYS)*100:.4f}%。'
        )

    code_step3 = (
        "# Step 3: Complete Portfolio (Utility Maximization)\n"
        "A = 3  # risk aversion coefficient\n"
        "# FOC: dU/dy = (mu_tang - rf) - A * y * sigma_tang^2 = 0\n"
        "y_star = (mu_tang - rf) / (A * sigma_tang**2)\n"
        "mu_complete = rf + y_star * (mu_tang - rf)\n"
        "sigma_complete = abs(y_star) * sigma_tang"
    )
    if lang == 'en':
        add_code_block(doc, code_step3, 'Key Code: Step 3')
    else:
        add_code_block(doc, code_step3, '核心代码：第三步')

    # ── Graphs ──
    if lang == 'en':
        doc.add_heading('5. Graphs', level=1)
        doc.add_paragraph(
            'Figure 1 shows the minimum-variance frontier (dashed = lower branch, solid = efficient frontier), '
            'the Capital Market Line (CML), the Global Minimum-Variance portfolio (green square), '
            'the tangency portfolio (red star), the complete portfolio (purple circle), '
            'and the 10 individual industry indices (black triangles).'
        )
    else:
        doc.add_heading('5. 图形分析', level=1)
        doc.add_paragraph(
            '图1展示了最小方差前沿（虚线为下半部分，实线为有效前沿）、'
            '资本市场线 (CML)、全局最小方差组合 (绿色方块)、'
            '切线组合 (红色五角星)、完整组合 (紫色圆圈) 以及10个行业指数 (黑色三角形)。'
        )

    doc.add_picture(ef_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if lang == 'en':
        doc.add_paragraph(
            'Figure 2 shows the tangency portfolio weights. '
            'Blue bars indicate long positions and red bars indicate short positions.'
        )
    else:
        doc.add_paragraph(
            '图2展示了切线组合的各行业权重。蓝色柱为多头持仓，红色柱为空头持仓。'
        )

    doc.add_picture(wt_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    suffix = '' if lang == 'en' else '_cn'
    doc_path = OUTPUT_DIR + f'task1_report{suffix}.docx'
    doc.save(doc_path)
    print(f"Saved {doc_path}")
    return doc_path


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("Task 1: Markowitz Portfolio Optimization")
    print("=" * 60)

    # ── Load data ──
    print("\n[1] Loading data...")
    returns = load_industry_data()
    rf_series = load_shibor(returns.index)
    rf_avg = rf_series.mean()
    rf_annual_avg_pct = rf_avg * 360 * 100
    print(f"    Returns shape: {returns.shape}")
    print(f"    Average daily rf: {rf_avg:.8f} ({rf_annual_avg_pct:.4f}% annualized)")

    # ── Descriptive statistics ──
    print("\n[2] Computing descriptive statistics...")
    stats = compute_descriptive_stats(returns)
    corr = compute_correlation(returns)
    print(stats.to_string())
    print(f"\nCorrelation matrix:\n{corr.to_string()}")

    # ── Markowitz computations ──
    mu = returns.mean().values
    Sigma = returns.cov().values

    eigvals = np.linalg.eigvals(Sigma)
    assert np.all(eigvals > 0), f"Covariance matrix not positive definite: {eigvals}"
    print(f"\n[3] Covariance matrix eigenvalues (all positive): min={eigvals.min():.2e}, max={eigvals.max():.2e}")

    # Step 2: Tangency Portfolio & CML (compute before frontier so we know the range)
    print("\n[5] Step 2: Computing tangency portfolio...")
    tangency = compute_tangency(mu, Sigma, rf_avg)
    print(f"    Tangency: E(r) = {tangency['mu']*TRADING_DAYS*100:.4f}% ann, "
          f"sigma = {tangency['sigma']*np.sqrt(TRADING_DAYS)*100:.4f}% ann, "
          f"Sharpe = {tangency['sharpe']*np.sqrt(TRADING_DAYS):.4f}")
    print(f"    Weights: {dict(zip([INDEX_NAMES[c] for c in INDEX_ORDER], [f'{w:.4f}' for w in tangency['w']]))}")

    # Step 1: Efficient Frontier (extend range to include tangency point)
    print("\n[4] Step 1: Computing efficient frontier...")
    frontier = compute_frontier_analytical(mu, Sigma, tang_mu=tangency['mu'])
    print(f"    GMV: E(r) = {frontier['mu_gmv']*TRADING_DAYS*100:.4f}% ann, "
          f"sigma = {frontier['sigma_gmv']*np.sqrt(TRADING_DAYS)*100:.4f}% ann")

    # Step 3: Complete Portfolio
    print("\n[6] Step 3: Computing complete portfolio...")
    y_star, mu_complete, sigma_complete = compute_complete_portfolio(
        tangency['mu'], tangency['sigma'], rf_avg, A_RISK_AVERSION)
    print(f"    y* = {y_star:.4f}")
    print(f"    Complete portfolio: E(r) = {mu_complete*TRADING_DAYS*100:.4f}% ann, "
          f"sigma = {sigma_complete*np.sqrt(TRADING_DAYS)*100:.4f}% ann")

    # ── Plotting ──
    print("\n[7] Generating plots...")
    for lang in ['en', 'cn']:
        ef_path = plot_efficient_frontier(frontier, tangency, rf_avg, mu, Sigma,
                                          y_star, mu_complete, sigma_complete, lang=lang)
        wt_path = plot_weights(tangency['w'], lang=lang)

    # ── Word document ──
    print("\n[8] Generating Word documents...")
    for lang in ['en', 'cn']:
        suffix = '' if lang == 'en' else '_cn'
        ef_path = OUTPUT_DIR + f'efficient_frontier{suffix}.png'
        wt_path = OUTPUT_DIR + f'portfolio_weights{suffix}.png'
        generate_report(stats, corr, frontier, tangency, rf_avg, rf_annual_avg_pct,
                        y_star, mu_complete, sigma_complete, returns, ef_path, wt_path, lang=lang)

    print("\n" + "=" * 60)
    print("Task 1 complete! Outputs in:", OUTPUT_DIR)
    print("=" * 60)


if __name__ == '__main__':
    main()

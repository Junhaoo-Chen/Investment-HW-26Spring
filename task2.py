#!/usr/bin/env python3
"""Task 2: Fama-French Two-Pass Regression for 10 SW Industry Indices."""

import numpy as np
import pandas as pd
import statsmodels.api as sm
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────────────────────────
DATA_DIR = '/cpfs02/data/shared/Group-m6/yanhai.cjh/Misc/investment_hw/数据文件/'
OUTPUT_DIR = '/cpfs02/data/shared/Group-m6/yanhai.cjh/Misc/investment_hw/output/'

INDEX_NAMES = {
    801010: '农林牧渔', 801050: '有色金属', 801080: '电子',
    801110: '家用电器', 801120: '食品饮料', 801150: '医药生物',
    801160: '公用事业', 801180: '房地产',   801780: '银行',
    801890: '机械设备',
}
INDEX_ORDER = [801010, 801050, 801080, 801110, 801120, 801150, 801160, 801180, 801780, 801890]


# ══════════════════════════════════════════════════════════════════════════════
# 1. DATA LOADING
# ══════════════════════════════════════════════════════════════════════════════

def load_industry_returns():
    files = ['20150101-20171231.xlsx', '20180101-20211231.xlsx', '20220101-20251231.xlsx']
    dfs = [pd.read_excel(DATA_DIR + f, skiprows=[1, 2]) for f in files]
    industry = pd.concat(dfs, ignore_index=True)
    returns_pct = industry.pivot(index='Idxtrd01', columns='Indexcd', values='Idxtrd08')
    returns_pct = returns_pct[INDEX_ORDER].sort_index()
    returns_decimal = returns_pct / 100.0
    return returns_decimal


def load_shibor(trading_dates):
    shibor = pd.read_excel(DATA_DIR + 'SHIBOR.xlsx', header=0)
    shibor = shibor.rename(columns={'交易日期_TrdDt': 'date', '加权价(%)_WghAvgPr': 'rate'})
    shibor = shibor[shibor['date'].isin(trading_dates)].set_index('date').sort_index()
    shibor['rf_daily'] = shibor['rate'] / (360 * 100)
    return shibor['rf_daily']


def load_ff3():
    df = pd.read_excel(DATA_DIR + '日频三因子.xlsx', skiprows=[1, 2])
    df = df[df['MarkettypeID'] == 'P9706'].copy()
    df = df.set_index('TradingDate').sort_index()
    return df[['RiskPremium1', 'SMB1', 'HML1']]


def load_ff5():
    df = pd.read_excel(DATA_DIR + '日频五因子.xlsx', skiprows=[1, 2])
    df = df[(df['MarkettypeID'] == 'P9706') & (df['Portfolios'] == 1)].copy()
    df = df.set_index('TradingDate').sort_index()
    return df[['RiskPremium1', 'SMB1', 'HML1', 'RMW1', 'CMA1']]


# ══════════════════════════════════════════════════════════════════════════════
# 2. FIRST-PASS: TIME-SERIES REGRESSIONS
# ══════════════════════════════════════════════════════════════════════════════

def first_pass_regression(excess_returns, factors, factor_names):
    X = sm.add_constant(factors)
    results = {}
    for col in excess_returns.columns:
        y = excess_returns[col]
        model = sm.OLS(y, X).fit()
        row = {'alpha': model.params['const'], 't_alpha': model.tvalues['const']}
        for fn in factor_names:
            row[f'beta_{fn}'] = model.params[fn]
            row[f't_{fn}'] = model.tvalues[fn]
        row['R2'] = model.rsquared
        row['Adj_R2'] = model.rsquared_adj
        results[col] = row
    return pd.DataFrame(results).T


def extract_betas(first_pass_df, factor_names):
    beta_cols = [f'beta_{fn}' for fn in factor_names]
    return first_pass_df[beta_cols].values


# ══════════════════════════════════════════════════════════════════════════════
# 3. SECOND-PASS: FAMA-MACBETH CROSS-SECTIONAL REGRESSIONS
# ══════════════════════════════════════════════════════════════════════════════

def second_pass_fama_macbeth(excess_returns, beta_hat):
    T = len(excess_returns)
    n_factors = beta_hat.shape[1]
    X_cross = sm.add_constant(beta_hat)

    gamma_ts = np.zeros((T, n_factors + 1))
    r2_ts = np.zeros(T)

    for t in range(T):
        y_cross = excess_returns.iloc[t].values
        model = sm.OLS(y_cross, X_cross).fit()
        gamma_ts[t] = model.params
        r2_ts[t] = model.rsquared

    gamma_mean = gamma_ts.mean(axis=0)
    gamma_std = gamma_ts.std(axis=0, ddof=1)
    gamma_tstat = gamma_mean / (gamma_std / np.sqrt(T))
    avg_r2 = r2_ts.mean()

    return gamma_mean, gamma_tstat, avg_r2, gamma_ts, r2_ts


# ══════════════════════════════════════════════════════════════════════════════
# 4. EXCEL OUTPUT
# ══════════════════════════════════════════════════════════════════════════════

def format_first_pass_table(df, factor_names):
    display_df = pd.DataFrame(index=[INDEX_NAMES[c] for c in df.index])
    display_df['alpha'] = df['alpha'].values
    display_df['t(alpha)'] = df['t_alpha'].values
    for fn in factor_names:
        display_df[f'beta_{fn}'] = df[f'beta_{fn}'].values
        display_df[f't({fn})'] = df[f't_{fn}'].values
    display_df['R^2'] = df['R2'].values
    display_df['Adj R^2'] = df['Adj_R2'].values
    return display_df


def format_second_pass_table(gamma_mean, gamma_tstat, avg_r2, labels):
    data = {
        'Statistic': ['Mean Estimate', 't-statistic'],
    }
    for i, lbl in enumerate(labels):
        data[lbl] = [gamma_mean[i], gamma_tstat[i]]
    data['Avg R^2'] = [avg_r2, '']
    return pd.DataFrame(data).set_index('Statistic')


def write_excel(excess_returns, ff3, ff5,
                fp_ff3, fp_ff5, sp_ff3_table, sp_ff5_table):
    path = OUTPUT_DIR + 'task2_results.xlsx'

    with pd.ExcelWriter(path, engine='xlsxwriter') as writer:
        workbook = writer.book

        header_fmt = workbook.add_format({
            'bold': True, 'bg_color': '#4472C4', 'font_color': 'white',
            'border': 1, 'text_wrap': True, 'align': 'center', 'valign': 'vcenter',
        })
        num_fmt = workbook.add_format({'num_format': '0.000000', 'border': 1})
        num_fmt4 = workbook.add_format({'num_format': '0.0000', 'border': 1})
        pct_fmt = workbook.add_format({'num_format': '0.00%', 'border': 1})

        # Sheet 1: Raw Data
        raw = pd.concat([
            excess_returns.rename(columns={c: f'ExRet_{INDEX_NAMES[c]}' for c in excess_returns.columns}),
            ff3.rename(columns={'RiskPremium1': 'MKT_RF', 'SMB1': 'SMB', 'HML1': 'HML'}),
            ff5[['RMW1', 'CMA1']].rename(columns={'RMW1': 'RMW', 'CMA1': 'CMA'}),
        ], axis=1)
        raw.to_excel(writer, sheet_name='Raw Data', index=True)
        ws = writer.sheets['Raw Data']
        ws.set_column('A:A', 12)
        ws.set_column('B:P', 14)

        # Sheet 2: FF3 First Pass
        fp_ff3.to_excel(writer, sheet_name='FF3 First Pass', index=True)
        ws = writer.sheets['FF3 First Pass']
        ws.set_column('A:A', 14)
        ws.set_column('B:K', 14)

        # Sheet 3: FF5 First Pass
        fp_ff5.to_excel(writer, sheet_name='FF5 First Pass', index=True)
        ws = writer.sheets['FF5 First Pass']
        ws.set_column('A:A', 14)
        ws.set_column('B:O', 14)

        # Sheet 4: FF3 Second Pass
        sp_ff3_table.to_excel(writer, sheet_name='FF3 Second Pass', index=True)
        ws = writer.sheets['FF3 Second Pass']
        ws.set_column('A:A', 16)
        ws.set_column('B:F', 16)

        # Sheet 5: FF5 Second Pass
        sp_ff5_table.to_excel(writer, sheet_name='FF5 Second Pass', index=True)
        ws = writer.sheets['FF5 Second Pass']
        ws.set_column('A:A', 16)
        ws.set_column('B:H', 16)

    print(f"Saved {path}")
    return path


# ══════════════════════════════════════════════════════════════════════════════
# 5. WORD ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_font(cell, text, font_name='Times New Roman', size=9, bold=False, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def add_df_table(doc, df, title, fmt='.4f', index_label=''):
    doc.add_paragraph(title, style='Heading 3')
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns) + 1)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_cell_font(table.rows[0].cells[0], index_label, bold=True, size=8)
    for j, col in enumerate(df.columns):
        set_cell_font(table.rows[0].cells[j + 1], str(col), bold=True, size=8)
    for i, (idx, row) in enumerate(df.iterrows()):
        set_cell_font(table.rows[i + 1].cells[0], str(idx), size=8)
        for j, val in enumerate(row):
            if isinstance(val, (float, np.floating)):
                set_cell_font(table.rows[i + 1].cells[j + 1], format(val, fmt), size=8,
                              alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                set_cell_font(table.rows[i + 1].cells[j + 1], str(val), size=8)


def generate_analysis_doc(fp_ff3, fp_ff5, sp_ff3_table, sp_ff5_table, lang='en'):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    if lang == 'en':
        doc.add_heading('Fama-French Factor Analysis Report', level=0)
        doc.add_paragraph('Project of Investment, 2026 — Task 2')

        # ── Section 1: Model Specification ──
        doc.add_heading('1. Regression Model Specification', level=1)

        doc.add_heading('1.1 Fama-French Three-Factor Model', level=2)
        doc.add_paragraph(
            'The Fama-French three-factor model explains asset excess returns using three risk factors: '
            'the market risk premium (MKT), the size factor (SMB, Small Minus Big), '
            'and the value factor (HML, High Minus Low book-to-market ratio).'
        )
        doc.add_paragraph(
            'Time-series regression (First Pass):',
        )
        doc.add_paragraph(
            '    R_i,t - R_f,t = alpha_i + beta_MKT,i * MKT_t + beta_SMB,i * SMB_t '
            '+ beta_HML,i * HML_t + epsilon_i,t',
            style='No Spacing'
        )
        doc.add_paragraph(
            'where R_i,t is the return of industry index i on day t, R_f,t is the daily risk-free rate '
            '(overnight SHIBOR / 36000), MKT_t = R_m,t - R_f,t is the market excess return, '
            'and SMB_t, HML_t are the size and value factors. All factor data are from CSMAR, '
            'using MarkettypeID = P9706 (all A-shares, free-float market-cap weighted).'
        )

        doc.add_heading('1.2 Fama-French Five-Factor Model', level=2)
        doc.add_paragraph(
            'The Fama-French five-factor model extends the three-factor model by adding two additional factors: '
            'the profitability factor (RMW, Robust Minus Weak) and the investment factor '
            '(CMA, Conservative Minus Aggressive).'
        )
        doc.add_paragraph(
            'Time-series regression (First Pass):',
        )
        doc.add_paragraph(
            '    R_i,t - R_f,t = alpha_i + beta_MKT,i * MKT_t + beta_SMB,i * SMB_t '
            '+ beta_HML,i * HML_t + beta_RMW,i * RMW_t + beta_CMA,i * CMA_t + epsilon_i,t',
            style='No Spacing'
        )

        doc.add_heading('1.3 Factor Construction', level=2)
        doc.add_paragraph(
            'All factors are constructed following Fama and French (2015) methodology using the 2x3 '
            'portfolio sorting approach (Portfolios = 1 in our data):'
        )
        doc.add_paragraph(
            'MKT (Market Risk Premium): The value-weighted return of all A-share stocks minus '
            'the risk-free rate (SHIBOR overnight). This is provided as RiskPremium1 in the CSMAR dataset.'
        )
        doc.add_paragraph(
            'SMB (Small Minus Big): Stocks are sorted into two groups by size (market capitalization) '
            'using the median as the breakpoint. SMB is the average return of the three small portfolios '
            'minus the average return of the three big portfolios, across the three B/M sorts.'
        )
        doc.add_paragraph(
            'HML (High Minus Low): Stocks are independently sorted into three groups by book-to-market '
            'ratio using the 30th and 70th percentile breakpoints. HML is the average return of the two '
            'high B/M portfolios minus the average return of the two low B/M portfolios.'
        )
        doc.add_paragraph(
            'RMW (Robust Minus Weak): Constructed analogously to HML, but sorting on operating '
            'profitability (revenues minus cost of goods sold, minus interest expense, minus SG&A, '
            'divided by book equity). RMW is the average return of the two robust (high profitability) '
            'portfolios minus the average return of the two weak (low profitability) portfolios, '
            'from independent 2x3 sorts on size and operating profitability.'
        )
        doc.add_paragraph(
            'CMA (Conservative Minus Aggressive): Constructed analogously, but sorting on investment '
            '(the growth rate of total assets). CMA is the average return of the two conservative '
            '(low investment) portfolios minus the average return of the two aggressive (high investment) '
            'portfolios, from independent 2x3 sorts on size and investment.'
        )

        doc.add_heading('1.4 Cross-Sectional Regression (Second Pass — Fama-MacBeth)', level=2)
        doc.add_paragraph(
            'For each trading day t, a cross-sectional regression is run across the 10 industry indices:'
        )
        doc.add_paragraph(
            '    R_i,t - R_f,t = gamma_0,t + gamma_1,t * beta_hat_MKT,i '
            '+ gamma_2,t * beta_hat_SMB,i + gamma_3,t * beta_hat_HML,i + eta_i,t',
            style='No Spacing'
        )
        doc.add_paragraph(
            'where beta_hat values are the estimated factor loadings from the first-pass time-series '
            'regressions (full-sample estimates). The Fama-MacBeth procedure averages the gamma '
            'coefficients across all T = 2,674 trading days:'
        )
        doc.add_paragraph(
            '    gamma_bar_k = (1/T) * sum(gamma_k,t)\n'
            '    SE(gamma_bar_k) = std(gamma_k,t) / sqrt(T)\n'
            '    t-statistic = gamma_bar_k / SE(gamma_bar_k)',
            style='No Spacing'
        )
        doc.add_paragraph(
            'The same procedure is applied for the five-factor model with 5 beta regressors. '
            'Note: with only 10 cross-sectional observations, the FF3 second pass has 6 degrees of freedom '
            '(10 - 4 parameters) and the FF5 second pass has 4 degrees of freedom (10 - 6 parameters). '
            'This is a limitation inherent in using industry-level (rather than individual stock or '
            'portfolio-level) test assets.'
        )

        # ── Section 2: First Pass Results ──
        doc.add_heading('2. First-Pass Results (Time-Series Regressions)', level=1)

        add_df_table(doc, fp_ff3, 'Table 1: FF3 First-Pass Regression Results', fmt='.4f', index_label='Industry')
        doc.add_paragraph('')

        doc.add_paragraph(
            'Interpretation of FF3 first-pass results:'
        )

        sig_alpha_ff3 = fp_ff3[fp_ff3['t(alpha)'].abs() > 1.96]
        if len(sig_alpha_ff3) > 0:
            names = ', '.join(sig_alpha_ff3.index)
            doc.add_paragraph(
                f'Alpha is statistically significant at the 5% level for: {names}. '
                'A significant alpha indicates that the three-factor model does not fully explain '
                'the average excess return of these industries.'
            )
        else:
            doc.add_paragraph(
                'No industry has a statistically significant alpha at the 5% level, suggesting '
                'that the three-factor model adequately explains the average excess returns.'
            )

        doc.add_paragraph(
            f'The market factor (MKT) loadings are all positive and highly significant across all '
            f'10 industries, confirming that market risk is the dominant driver of industry returns. '
            f'R-squared values range from {fp_ff3["R^2"].min():.4f} to {fp_ff3["R^2"].max():.4f}, '
            f'with an average of {fp_ff3["R^2"].mean():.4f}.'
        )

        doc.add_paragraph('')
        add_df_table(doc, fp_ff5, 'Table 2: FF5 First-Pass Regression Results', fmt='.4f', index_label='Industry')
        doc.add_paragraph('')

        doc.add_paragraph(
            'Interpretation of FF5 first-pass results:'
        )

        sig_alpha_ff5 = fp_ff5[fp_ff5['t(alpha)'].abs() > 1.96]
        if len(sig_alpha_ff5) > 0:
            names = ', '.join(sig_alpha_ff5.index)
            doc.add_paragraph(
                f'Alpha is statistically significant at the 5% level for: {names}.'
            )
        else:
            doc.add_paragraph(
                'No industry has a statistically significant alpha at the 5% level under the five-factor model.'
            )

        doc.add_paragraph(
            f'Compared to the three-factor model, the five-factor model generally provides '
            f'higher R-squared values (average {fp_ff5["R^2"].mean():.4f} vs {fp_ff3["R^2"].mean():.4f}), '
            f'indicating improved explanatory power. The additional factors RMW and CMA capture '
            f'variation in industry returns that is not explained by MKT, SMB, and HML alone.'
        )

        # ── Section 3: Second Pass Results ──
        doc.add_heading('3. Second-Pass Results (Fama-MacBeth Cross-Sectional Regressions)', level=1)

        add_df_table(doc, sp_ff3_table, 'Table 3: FF3 Fama-MacBeth Second-Pass Results', fmt='.4f',
                     index_label='Statistic')
        doc.add_paragraph('')

        gamma0_ff3 = sp_ff3_table.iloc[0, 0] if isinstance(sp_ff3_table.iloc[0, 0], float) else 0
        t_gamma0_ff3 = sp_ff3_table.iloc[1, 0] if isinstance(sp_ff3_table.iloc[1, 0], float) else 0

        doc.add_paragraph(
            f'The intercept (gamma_0) has a mean estimate of {gamma0_ff3:.6f} with a t-statistic of '
            f'{t_gamma0_ff3:.4f}. '
            + ('This is statistically significant, suggesting the FF3 model does not fully capture '
               'the cross-sectional variation in expected returns — there exists mispricing or missing factors.'
               if abs(t_gamma0_ff3) > 1.96 else
               'This is not statistically significant, consistent with the FF3 model adequately explaining '
               'the cross-sectional variation in expected returns.')
        )

        doc.add_paragraph('')
        add_df_table(doc, sp_ff5_table, 'Table 4: FF5 Fama-MacBeth Second-Pass Results', fmt='.4f',
                     index_label='Statistic')
        doc.add_paragraph('')

        gamma0_ff5 = sp_ff5_table.iloc[0, 0] if isinstance(sp_ff5_table.iloc[0, 0], float) else 0
        t_gamma0_ff5 = sp_ff5_table.iloc[1, 0] if isinstance(sp_ff5_table.iloc[1, 0], float) else 0

        doc.add_paragraph(
            f'Under the five-factor model, the intercept (gamma_0) has a mean estimate of {gamma0_ff5:.6f} '
            f'with a t-statistic of {t_gamma0_ff5:.4f}. '
            + ('This remains statistically significant.'
               if abs(t_gamma0_ff5) > 1.96 else
               'This is not statistically significant, supporting the FF5 model.')
        )

        # ── Section 4: Conclusion ──
        doc.add_heading('4. Conclusion', level=1)
        doc.add_paragraph(
            'This study applies the Fama-MacBeth two-pass regression methodology to test whether '
            'the Fama-French three-factor and five-factor models can explain the equity premium '
            'of 10 Shenwan Level-1 industry indices in the Chinese A-share market over 2015-2025. '
            'The first-pass time-series regressions show that market risk (MKT) is the dominant factor, '
            'with high and significant loadings across all industries. The size (SMB) and value (HML) '
            'factors provide additional explanatory power, and the profitability (RMW) and investment (CMA) '
            'factors in the five-factor model further improve R-squared. '
            'The second-pass Fama-MacBeth cross-sectional regressions test whether these factor loadings '
            'are priced in the cross-section. A caveat is that with only 10 test assets (industry indices), '
            'the cross-sectional regressions have limited degrees of freedom, which may reduce the power '
            'of the tests.'
        )

    else:
        doc.add_heading('Fama-French 因子分析报告', level=0)
        doc.add_paragraph('2026年投资学课程项目 — 任务二')

        doc.add_heading('1. 回归模型设定', level=1)

        doc.add_heading('1.1 Fama-French 三因子模型', level=2)
        doc.add_paragraph(
            'Fama-French 三因子模型使用三个风险因子解释资产超额收益：'
            '市场风险溢价 (MKT)、规模因子 (SMB, Small Minus Big) 和价值因子 (HML, High Minus Low)。'
        )
        doc.add_paragraph('时间序列回归（第一步）：')
        doc.add_paragraph(
            '    R_i,t - R_f,t = alpha_i + beta_MKT,i * MKT_t + beta_SMB,i * SMB_t '
            '+ beta_HML,i * HML_t + epsilon_i,t',
            style='No Spacing'
        )
        doc.add_paragraph(
            '其中 R_i,t 为行业指数 i 在 t 日的收益率，R_f,t 为日度无风险利率'
            '（SHIBOR隔夜利率 / 36000），MKT_t = R_m,t - R_f,t 为市场超额收益，'
            'SMB_t、HML_t 分别为规模因子和价值因子。因子数据来源于国泰安 (CSMAR) 数据库，'
            '使用 MarkettypeID = P9706（综合A股，流通市值加权）。'
        )

        doc.add_heading('1.2 Fama-French 五因子模型', level=2)
        doc.add_paragraph(
            'Fama-French 五因子模型在三因子基础上增加了盈利能力因子 (RMW, Robust Minus Weak) '
            '和投资模式因子 (CMA, Conservative Minus Aggressive)。'
        )
        doc.add_paragraph('时间序列回归（第一步）：')
        doc.add_paragraph(
            '    R_i,t - R_f,t = alpha_i + beta_MKT,i * MKT_t + beta_SMB,i * SMB_t '
            '+ beta_HML,i * HML_t + beta_RMW,i * RMW_t + beta_CMA,i * CMA_t + epsilon_i,t',
            style='No Spacing'
        )

        doc.add_heading('1.3 因子构建方法', level=2)
        doc.add_paragraph(
            '所有因子按照 Fama and French (2015) 方法，采用 2x3 组合排序方法构建（数据中 Portfolios = 1）：'
        )
        doc.add_paragraph(
            'MKT（市场风险溢价因子）：全部A股流通市值加权收益率减去无风险利率（SHIBOR隔夜利率）。'
            '在 CSMAR 数据中以 RiskPremium1 提供。'
        )
        doc.add_paragraph(
            'SMB（规模因子）：按市值中位数将股票分为大、小两组，SMB 等于三个小市值组合的平均收益率'
            '减去三个大市值组合的平均收益率。'
        )
        doc.add_paragraph(
            'HML（价值因子）：按账面市值比的第30和第70百分位将股票分为低、中、高三组。'
            'HML 等于两个高账面市值比组合的平均收益率减去两个低账面市值比组合的平均收益率。'
        )
        doc.add_paragraph(
            'RMW（盈利能力因子）：按营业利润率（营业收入减去成本、利息和销售管理费用，除以账面价值）'
            '将股票分组。RMW 等于两个高盈利组合的平均收益率减去两个低盈利组合的平均收益率。'
        )
        doc.add_paragraph(
            'CMA（投资模式因子）：按总资产增长率将股票分组。CMA 等于两个保守（低投资）组合的'
            '平均收益率减去两个激进（高投资）组合的平均收益率。'
        )

        doc.add_heading('1.4 截面回归（第二步 — Fama-MacBeth 方法）', level=2)
        doc.add_paragraph(
            '对每个交易日 t，对10个行业指数进行截面回归：'
        )
        doc.add_paragraph(
            '    R_i,t - R_f,t = gamma_0,t + gamma_1,t * beta_hat_MKT,i '
            '+ gamma_2,t * beta_hat_SMB,i + gamma_3,t * beta_hat_HML,i + eta_i,t',
            style='No Spacing'
        )
        doc.add_paragraph(
            '其中 beta_hat 为第一步时间序列回归中估计的全样本因子载荷。Fama-MacBeth 方法对 T = 2,674 '
            '个交易日的 gamma 系数求均值：'
        )
        doc.add_paragraph(
            '    gamma_bar_k = (1/T) * sum(gamma_k,t)\n'
            '    SE(gamma_bar_k) = std(gamma_k,t) / sqrt(T)\n'
            '    t-statistic = gamma_bar_k / SE(gamma_bar_k)',
            style='No Spacing'
        )
        doc.add_paragraph(
            '五因子模型的截面回归类似，使用5个 beta 作为解释变量。'
            '注意：由于仅有10个截面观测（行业指数），三因子模型的截面回归有6个自由度'
            '（10 - 4个参数），五因子模型仅有4个自由度（10 - 6个参数），'
            '这是使用行业级（而非个股或组合级）检验资产的固有局限。'
        )

        doc.add_heading('2. 第一步回归结果（时间序列回归）', level=1)
        add_df_table(doc, fp_ff3, 'Table 1: 三因子模型第一步回归结果', fmt='.4f', index_label='行业')
        doc.add_paragraph('')

        sig_alpha_ff3 = fp_ff3[fp_ff3['t(alpha)'].abs() > 1.96]
        if len(sig_alpha_ff3) > 0:
            names = ', '.join(sig_alpha_ff3.index)
            doc.add_paragraph(
                f'在5%显著性水平下，以下行业的 alpha 显著不为零：{names}。'
                '显著的 alpha 表明三因子模型未能完全解释这些行业的平均超额收益。'
            )
        else:
            doc.add_paragraph(
                '在5%显著性水平下，没有行业的 alpha 显著不为零，'
                '说明三因子模型能够较好地解释各行业的平均超额收益。'
            )

        doc.add_paragraph(
            f'市场因子 (MKT) 的载荷在所有10个行业中均为正且高度显著，'
            f'表明市场风险是行业收益的主要驱动力。'
            f'R方值范围为 {fp_ff3["R^2"].min():.4f} 至 {fp_ff3["R^2"].max():.4f}，'
            f'平均为 {fp_ff3["R^2"].mean():.4f}。'
        )

        doc.add_paragraph('')
        add_df_table(doc, fp_ff5, 'Table 2: 五因子模型第一步回归结果', fmt='.4f', index_label='行业')
        doc.add_paragraph('')

        sig_alpha_ff5 = fp_ff5[fp_ff5['t(alpha)'].abs() > 1.96]
        if len(sig_alpha_ff5) > 0:
            names = ', '.join(sig_alpha_ff5.index)
            doc.add_paragraph(f'在5%显著性水平下，以下行业的 alpha 显著不为零：{names}。')
        else:
            doc.add_paragraph('在五因子模型下，没有行业的 alpha 在5%水平上显著不为零。')

        doc.add_paragraph(
            f'与三因子模型相比，五因子模型的R方值普遍更高'
            f'（平均 {fp_ff5["R^2"].mean():.4f} vs {fp_ff3["R^2"].mean():.4f}），'
            f'表明盈利能力因子 (RMW) 和投资模式因子 (CMA) 捕获了三因子无法解释的行业收益变异。'
        )

        doc.add_heading('3. 第二步回归结果（Fama-MacBeth 截面回归）', level=1)
        add_df_table(doc, sp_ff3_table, 'Table 3: 三因子 Fama-MacBeth 第二步回归结果', fmt='.4f',
                     index_label='统计量')
        doc.add_paragraph('')

        gamma0_ff3 = sp_ff3_table.iloc[0, 0] if isinstance(sp_ff3_table.iloc[0, 0], float) else 0
        t_gamma0_ff3 = sp_ff3_table.iloc[1, 0] if isinstance(sp_ff3_table.iloc[1, 0], float) else 0

        doc.add_paragraph(
            f'截距项 (gamma_0) 的均值估计为 {gamma0_ff3:.6f}，t 统计量为 {t_gamma0_ff3:.4f}。'
            + ('该截距显著不为零，说明三因子模型未能完全解释截面收益差异，可能存在定价误差或遗漏因子。'
               if abs(t_gamma0_ff3) > 1.96 else
               '该截距不显著，与三因子模型充分解释截面收益差异的假说一致。')
        )

        doc.add_paragraph('')
        add_df_table(doc, sp_ff5_table, 'Table 4: 五因子 Fama-MacBeth 第二步回归结果', fmt='.4f',
                     index_label='统计量')
        doc.add_paragraph('')

        gamma0_ff5 = sp_ff5_table.iloc[0, 0] if isinstance(sp_ff5_table.iloc[0, 0], float) else 0
        t_gamma0_ff5 = sp_ff5_table.iloc[1, 0] if isinstance(sp_ff5_table.iloc[1, 0], float) else 0

        doc.add_paragraph(
            f'五因子模型下，截距项 (gamma_0) 的均值估计为 {gamma0_ff5:.6f}，'
            f't 统计量为 {t_gamma0_ff5:.4f}。'
            + ('该截距仍然显著。' if abs(t_gamma0_ff5) > 1.96 else '该截距不显著，支持五因子模型。')
        )

        doc.add_heading('4. 结论', level=1)
        doc.add_paragraph(
            '本研究采用 Fama-MacBeth 两步回归法，检验 Fama-French 三因子和五因子模型能否解释'
            '2015-2025年中国A股市场10个申万一级行业指数的股权溢价。'
            '第一步时间序列回归表明，市场风险 (MKT) 是最主要的解释因子，在所有行业中均有高度显著的载荷。'
            '规模 (SMB) 和价值 (HML) 因子提供了额外的解释力，而五因子模型中的盈利能力 (RMW) 和'
            '投资模式 (CMA) 因子进一步提高了R方。'
            '第二步 Fama-MacBeth 截面回归检验了这些因子载荷在截面上是否被定价。'
            '需要注意的是，由于仅有10个检验资产（行业指数），截面回归的自由度有限，可能降低检验功效。'
        )

    suffix = '' if lang == 'en' else '_cn'
    doc_path = OUTPUT_DIR + f'task2_analysis{suffix}.docx'
    doc.save(doc_path)
    print(f"Saved {doc_path}")
    return doc_path


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("Task 2: Fama-French Two-Pass Regression")
    print("=" * 60)

    # ── Load data ──
    print("\n[1] Loading data...")
    returns = load_industry_returns()
    rf = load_shibor(returns.index)
    ff3 = load_ff3()
    ff5 = load_ff5()

    common_dates = returns.index.intersection(ff3.index).intersection(ff5.index).intersection(rf.index)
    returns = returns.loc[common_dates]
    rf = rf.loc[common_dates]
    ff3 = ff3.loc[common_dates]
    ff5 = ff5.loc[common_dates]

    excess_returns = returns.sub(rf, axis=0)

    print(f"    Excess returns shape: {excess_returns.shape}")
    print(f"    FF3 shape: {ff3.shape}")
    print(f"    FF5 shape: {ff5.shape}")
    assert excess_returns.shape[0] == ff3.shape[0] == ff5.shape[0]
    assert excess_returns.isna().sum().sum() == 0
    assert ff3.isna().sum().sum() == 0
    assert ff5.isna().sum().sum() == 0

    # ── First Pass ──
    print("\n[2] Running first-pass time-series regressions...")
    ff3_names = ['RiskPremium1', 'SMB1', 'HML1']
    ff5_names = ['RiskPremium1', 'SMB1', 'HML1', 'RMW1', 'CMA1']

    fp_ff3_raw = first_pass_regression(excess_returns, ff3[ff3_names], ff3_names)
    fp_ff5_raw = first_pass_regression(excess_returns, ff5[ff5_names], ff5_names)

    fp_ff3_display = format_first_pass_table(fp_ff3_raw, ff3_names)
    fp_ff5_display = format_first_pass_table(fp_ff5_raw, ff5_names)

    print("\n    FF3 First-Pass Results:")
    print(fp_ff3_display.to_string())
    print(f"\n    FF3 avg R^2: {fp_ff3_display['R^2'].mean():.4f}")

    print("\n    FF5 First-Pass Results:")
    print(fp_ff5_display.to_string())
    print(f"\n    FF5 avg R^2: {fp_ff5_display['R^2'].mean():.4f}")

    # Verify FF5 R^2 >= FF3 R^2
    for i, idx in enumerate(fp_ff3_display.index):
        r2_3 = fp_ff3_display.loc[idx, 'R^2']
        r2_5 = fp_ff5_display.loc[idx, 'R^2']
        assert r2_5 >= r2_3 - 1e-10, f"FF5 R^2 < FF3 R^2 for {idx}: {r2_5} < {r2_3}"

    # ── Second Pass ──
    print("\n[3] Running second-pass Fama-MacBeth regressions...")
    beta_ff3 = extract_betas(fp_ff3_raw, ff3_names)
    beta_ff5 = extract_betas(fp_ff5_raw, ff5_names)

    gm3, gt3, ar3, _, _ = second_pass_fama_macbeth(excess_returns, beta_ff3)
    gm5, gt5, ar5, _, _ = second_pass_fama_macbeth(excess_returns, beta_ff5)

    sp_ff3_labels = ['gamma_0 (intercept)', 'gamma_MKT', 'gamma_SMB', 'gamma_HML']
    sp_ff5_labels = ['gamma_0 (intercept)', 'gamma_MKT', 'gamma_SMB', 'gamma_HML', 'gamma_RMW', 'gamma_CMA']

    sp_ff3_table = format_second_pass_table(gm3, gt3, ar3, sp_ff3_labels)
    sp_ff5_table = format_second_pass_table(gm5, gt5, ar5, sp_ff5_labels)

    print("\n    FF3 Fama-MacBeth Results:")
    print(sp_ff3_table.to_string())
    print(f"\n    FF5 Fama-MacBeth Results:")
    print(sp_ff5_table.to_string())

    # ── Excel ──
    print("\n[4] Writing Excel output...")
    write_excel(excess_returns, ff3, ff5, fp_ff3_display, fp_ff5_display, sp_ff3_table, sp_ff5_table)

    # ── Word analysis ──
    print("\n[5] Generating Word analysis documents...")
    for lang in ['en', 'cn']:
        generate_analysis_doc(fp_ff3_display, fp_ff5_display, sp_ff3_table, sp_ff5_table, lang=lang)

    print("\n" + "=" * 60)
    print("Task 2 complete! Outputs in:", OUTPUT_DIR)
    print("=" * 60)


if __name__ == '__main__':
    main()
